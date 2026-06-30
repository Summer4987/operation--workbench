from pathlib import Path
import zipfile
import shutil
import tempfile
import subprocess

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from lxml import etree
from PIL import Image, ImageChops


WORK = Path(__file__).resolve().parent
SRC = Path("/Users/summer/Downloads/「熊小小牛排饭」门店全流程培训手册.docx")
PDF_SRC = Path("/Users/summer/Library/Containers/com.tencent.xinWeChat/Data/Documents/xwechat_files/wxid_7tw4mnbss0yk12_c37e/temp/drag/「熊小小牛排饭」门店全流程培训.pdf")
OUT = WORK / "「熊小小牛排饭」门店全流程培训手册_重排版.docx"
MEDIA_DIR = WORK / "manual_reformat_media"
PDF_RENDER_DIR = WORK / "manual_pdf_production_pages"
FONT = "PingFang SC"


def extract_media():
    if MEDIA_DIR.exists():
        shutil.rmtree(MEDIA_DIR)
    MEDIA_DIR.mkdir(parents=True)
    with zipfile.ZipFile(SRC) as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                (MEDIA_DIR / Path(name).name).write_bytes(z.read(name))


def crop_whitespace(path, padding=18):
    image = Image.open(path).convert("RGB")
    bg = Image.new("RGB", image.size, "white")
    diff = Image.eval(ImageChops.difference(image, bg), lambda px: 255 if px > 18 else 0)
    box = diff.getbbox()
    if not box:
        return path
    left, top, right, bottom = box
    left = max(0, left - padding)
    top = max(0, top - padding)
    right = min(image.width, right + padding)
    bottom = min(image.height, bottom + padding)
    out = path.with_name(path.stem + "_cropped.png")
    image.crop((left, top, right, bottom)).save(out)
    return out


def render_pdf_production_pages():
    if PDF_RENDER_DIR.exists():
        shutil.rmtree(PDF_RENDER_DIR)
    PDF_RENDER_DIR.mkdir(parents=True)
    fallback_dir = WORK / "pdf_pages"
    if not PDF_SRC.exists():
        pages = []
        for page in range(11, 22):
            raw = fallback_dir / f"page-{page}.png"
            if raw.exists():
                copied = PDF_RENDER_DIR / f"production-{page}.png"
                shutil.copy(raw, copied)
                pages.append((page, crop_whitespace(copied)))
        return pages
    pdftoppm = Path("/Users/summer/.cache/codex-runtimes/codex-primary-runtime/dependencies/bin/pdftoppm")
    subprocess.run(
        [str(pdftoppm), "-png", "-r", "180", "-f", "11", "-l", "21", str(PDF_SRC), str(PDF_RENDER_DIR / "production")],
        check=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
    )
    pages = []
    for page in range(11, 22):
        raw = PDF_RENDER_DIR / f"production-{page}.png"
        if raw.exists():
            pages.append((page, crop_whitespace(raw)))
    return pages


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=9.5, color="222222"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.bold = bold
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(section):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def new_page(doc):
    doc.add_page_break()


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.1)
    section.right_margin = Cm(1.1)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(34, 34, 34)
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color, before, after in [
        ("Title", 24, "1F4E5F", 0, 12),
        ("Heading 1", 16, "1F4E5F", 14, 6),
        ("Heading 2", 12.5, "3C6E71", 8, 4),
        ("Heading 3", 10.5, "4D4D4D", 5, 2),
    ]:
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = False


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(86)
    r = p.add_run("「熊小小牛排饭」\n门店全流程培训手册")
    r.bold = True
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    r.font.size = Pt(30)
    r.font.color.rgb = RGBColor.from_string("1F4E5F")

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_before = Pt(12)
    rr = sub.add_run("设备 · 流程 · 出餐 · 食安 · 客诉")
    rr.font.name = FONT
    rr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    rr._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    rr._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    rr.font.size = Pt(13)
    rr.font.color.rgb = RGBColor.from_string("666666")

    deco = doc.add_paragraph()
    deco.alignment = WD_ALIGN_PARAGRAPH.CENTER
    deco.paragraph_format.space_before = Pt(22)
    dr = deco.add_run("门店操作标准 · 新员工培训 · 日常复盘")
    dr.font.name = FONT
    dr._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    dr._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    dr._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    dr.font.size = Pt(11)
    dr.font.color.rgb = RGBColor.from_string("3C6E71")

    bottom = doc.add_paragraph()
    bottom.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom.paragraph_format.space_before = Pt(180)
    br = bottom.add_run("内部培训手册")
    br.font.name = FONT
    br._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    br._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    br._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    br.font.size = Pt(10)
    br.font.color.rgb = RGBColor.from_string("777777")
    new_page(doc)


def add_contents(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("目录")
    r.bold = True
    r.font.name = FONT
    r._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    r._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    r._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    r.font.size = Pt(22)
    r.font.color.rgb = RGBColor.from_string("1F4E5F")
    items = [
        ("一", "基础设备使用和维护", "3"),
        ("二", "全天工作流程", "7"),
        ("三", "餐品制作和打包", "9"),
        ("四", "食品安全", "19"),
        ("五", "每日闭店拍照", "22"),
        ("六", "异常订单处理方式", "23"),
        ("七", "客诉参考处理方式", "24"),
    ]
    table = doc.add_table(rows=len(items), cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, (num, title, page) in enumerate(items):
        set_cell_text(table.cell(idx, 0), num, bold=True, size=12, color="1F4E5F")
        set_cell_text(table.cell(idx, 1), title, size=12)
        set_cell_text(table.cell(idx, 2), page, bold=True, size=12, color="1F4E5F")
        table.cell(idx, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(idx, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(idx, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table.cell(idx, 0).width = Cm(1.2)
        table.cell(idx, 1).width = Cm(12.2)
        table.cell(idx, 2).width = Cm(1.4)
    note = doc.add_paragraph("阅读建议：新员工先通读一、二、四、五章；上岗前重点复习第三章图表；遇到售后问题时直接翻第六、七章。")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    new_page(doc)


def add_para(doc, text, style=None, bold_lead=False):
    p = doc.add_paragraph(style=style)
    if bold_lead and "：" in text:
        lead, rest = text.split("：", 1)
        r = p.add_run(lead + "：")
        r.bold = True
        p.add_run(rest)
    else:
        p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        add_para(doc, item, style="List Bullet", bold_lead=True)


def add_numbers(doc, items):
    for item in items:
        add_para(doc, item, style="List Number")


def resolve_image_path(filename):
    path = MEDIA_DIR / filename
    if not path.exists():
        return None
    if path.suffix.lower() in [".jpg", ".jpeg"]:
        converted = path.with_suffix(".png")
        if not converted.exists():
            Image.open(path).convert("RGB").save(converted)
        path = converted
    return path


def add_picture_fit(run, path, max_width_cm=15.8, max_height_cm=None):
    image = Image.open(path)
    width_px, height_px = image.size
    ratio = width_px / height_px
    if max_height_cm is None:
        run.add_picture(str(path), width=Cm(max_width_cm))
        return
    width_if_max_height = max_height_cm * ratio
    if width_if_max_height <= max_width_cm:
        run.add_picture(str(path), height=Cm(max_height_cm))
    else:
        run.add_picture(str(path), width=Cm(max_width_cm))


def add_figure(doc, filename, caption, width_cm=15.8, height_cm=None):
    path = resolve_image_path(filename)
    if path is None:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture_fit(p.add_run(), path, max_width_cm=width_cm, max_height_cm=height_cm)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = RGBColor.from_string("666666")


def add_equipment_card(doc, name, use, image, notes, image_width=6.4, image_height=5.0):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    left, right = table.rows[0].cells
    set_cell_shading(left, "F7FAF9")
    set_cell_shading(right, "FFFFFF")
    left.width = Cm(11.2)
    right.width = Cm(4.7)

    left.text = ""
    p = left.paragraphs[0]
    title = p.add_run(f"{name}：{use}")
    title.bold = True
    title.font.size = Pt(10.5)
    title.font.color.rgb = RGBColor.from_string("1F4E5F")
    for note in notes:
        bp = left.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_after = Pt(1)
        bp.add_run(note)

    right.text = ""
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = resolve_image_path(image)
    if path:
        add_picture_fit(rp.add_run(), path, max_width_cm=image_width, max_height_cm=image_height)
    for cell in [left, right]:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_key_table(doc, rows, widths=(2.7, 12.5)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], "E8F1EF")
        set_cell_text(cells[0], label, bold=True, color="1F4E5F")
        set_cell_text(cells[1], value)
        cells[0].width = Cm(widths[0])
        cells[1].width = Cm(widths[1])
    return table


def add_figure_grid(doc, entries, columns=2):
    rows = (len(entries) + columns - 1) // columns
    table = doc.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, entry in enumerate(entries):
        title, image, caption, max_w, max_h = entry
        cell = table.cell(idx // columns, idx % columns)
        cell.text = ""
        hp = cell.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hr = hp.add_run(title)
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor.from_string("1F4E5F")
        path = resolve_image_path(image)
        if path:
            ip = cell.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_picture_fit(ip.add_run(), path, max_width_cm=max_w, max_height_cm=max_h)
        cp = cell.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].font.size = Pt(8)
        cp.runs[0].font.color.rgb = RGBColor.from_string("666666")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def split_image(filename, parts, prefix):
    path = resolve_image_path(filename)
    if path is None:
        return []
    crop_dir = MEDIA_DIR / "crops"
    crop_dir.mkdir(exist_ok=True)
    image = Image.open(path).convert("RGB")
    width, height = image.size
    chunk_h = height // parts
    outputs = []
    overlap = min(60, max(0, height // 80))
    for idx in range(parts):
        top = max(0, idx * chunk_h - (overlap if idx else 0))
        bottom = height if idx == parts - 1 else min(height, (idx + 1) * chunk_h + overlap)
        cropped = image.crop((0, top, width, bottom))
        out = crop_dir / f"{prefix}_{idx + 1}.png"
        cropped.save(out)
        outputs.append(out)
    return outputs


def split_image_by_ranges(filename, ranges, prefix):
    path = resolve_image_path(filename)
    if path is None:
        return []
    crop_dir = MEDIA_DIR / "crops"
    crop_dir.mkdir(exist_ok=True)
    image = Image.open(path).convert("RGB")
    width, height = image.size
    outputs = []
    for idx, (top, bottom) in enumerate(ranges, 1):
        top = max(0, int(top))
        bottom = min(height, int(bottom))
        cropped = image.crop((0, top, width, bottom))
        out = crop_dir / f"{prefix}_{idx}.png"
        cropped.save(out)
        outputs.append(out)
    return outputs


def add_split_figure(doc, filename, title, caption, parts, width_cm=15.8):
    doc.add_heading(title, level=2)
    chunks = split_image(filename, parts, title)
    for idx, path in enumerate(chunks, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture_fit(p.add_run(), path, max_width_cm=width_cm)
        cap = doc.add_paragraph(f"{caption}（{idx}/{len(chunks)}）")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor.from_string("666666")
        if idx != len(chunks):
            new_page(doc)


def add_ranged_split_figure(doc, filename, title, caption, ranges, width_cm=15.8):
    doc.add_heading(title, level=2)
    chunks = split_image_by_ranges(filename, ranges, title)
    for idx, path in enumerate(chunks, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture_fit(p.add_run(), path, max_width_cm=width_cm)
        cap = doc.add_paragraph(f"{caption}（{idx}/{len(chunks)}）")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor.from_string("666666")
        if idx != len(chunks):
            new_page(doc)


def add_ranged_images_compact(doc, filename, title, ranges, width_cm=15.2):
    doc.add_heading(title, level=2)
    chunks = split_image_by_ranges(filename, ranges, title)
    for idx, path in enumerate(chunks, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture_fit(p.add_run(), path, max_width_cm=width_cm)
        if idx != len(chunks):
            new_page(doc)


def make_packed_item_pages(filename, boundaries, prefix, width_cm, max_display_height_cm):
    path = resolve_image_path(filename)
    if path is None:
        return []
    crop_dir = MEDIA_DIR / "packed"
    crop_dir.mkdir(exist_ok=True)
    image = Image.open(path).convert("RGB")
    native_max_height = int(max_display_height_cm / width_cm * image.width)
    blocks = []
    for idx, (top, bottom) in enumerate(zip(boundaries[:-1], boundaries[1:]), 1):
        top = max(0, int(top))
        bottom = min(image.height, int(bottom))
        if bottom <= top:
            continue
        blocks.append((idx, top, bottom, image.crop((0, top, image.width, bottom))))

    pages = []
    current = []
    current_h = 0
    for block in blocks:
        _, _, _, block_img = block
        block_h = block_img.height
        if current and current_h + block_h > native_max_height:
            pages.append(current)
            current = []
            current_h = 0
        current.append(block)
        current_h += block_h
    if current:
        pages.append(current)

    outputs = []
    for page_idx, page_blocks in enumerate(pages, 1):
        total_h = sum(block_img.height for _, _, _, block_img in page_blocks)
        canvas = Image.new("RGB", (image.width, total_h), "white")
        y = 0
        for _, _, _, block_img in page_blocks:
            canvas.paste(block_img, (0, y))
            y += block_img.height
        out = crop_dir / f"{prefix}_{page_idx}.png"
        canvas.save(out)
        outputs.append(out)
    return outputs


def add_packed_item_images(doc, filename, title, boundaries, width_cm=14.0, max_display_height_cm=18.6):
    doc.add_heading(title, level=2)
    pages = make_packed_item_pages(filename, boundaries, title, width_cm, max_display_height_cm)
    for idx, path in enumerate(pages, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture_fit(p.add_run(), path, max_width_cm=width_cm)
        if idx != len(pages):
            new_page(doc)


def chapter_equipment(doc):
    doc.add_heading("一、基础设备使用和维护", level=1)
    doc.add_heading("了解基础设备使用方式", level=2)
    equipment = [
        ("烤炉", "用于烤牛排（牛五花、板腱、眼肉、菲力）", "image1.png", [
            "整个烤炉分三个区，每个开关分别控制左中右三部分，有1-10个按键式档位，平时烤肉必须开到10档。",
            "从开机上温到完全达到温度需要3-5分钟，高峰期烤炉不关，平峰期待机时可调至3-5档，下午非出餐时段可以关。",
            "底部有接油盘，烤炉每天使用前需要加水，水烧干需要添加水，切勿干烧。",
            "牛排放在烤网上在上炉烤制，不可以直接放在烤架上烤。",
        ]),
        ("电磁炉", "用于炒蔬菜、煮蔬菜、炒嫩肩牛肉、烫藤椒牛肉", "image2.png", [
            "不同品牌的电磁炉最大档位设定不同，炒、煮蔬菜和肉类时全部用最大火。",
            "煮土豆、煮全熟鸡蛋用中火，化黄油用中小火。",
        ]),
        ("烤箱", "用于烤三文鱼、烤胡萝卜、烤手枪腿", "image3.png", [
            "开机前务必确保烤箱内清空，无烤盘、食物、其他物品。",
            "长按开机；首次设置温度为200度，后续温度默认200度。",
            "按显示屏第三行旁边的上下按键调节烤制时间，再按左侧第一列第二个按键开始工作。",
            "预热完成并显示「FOOD」后，迅速放入已摆好食材的烤盘，关门后时间自动倒计时。",
            "倒计时结束后立刻取出食材；厨房嘈杂时另设定时器，避免烤过火。",
        ]),
        ("电饭锅", "用于煮米饭", "image4.png", [
            "常用模式为精煮饭和快煮饭，精煮饭约50分钟，快煮饭约30分钟；精煮饭效果更好。",
            "屏幕显示E且按键无反应时，长按取消键可恢复正常。",
            "预约功能显示的是多少小时后煮好饭，不是开始煮饭的时间。",
        ]),
        ("微波炉", "用于快速加热", "image5.png", [
            "购买时请卖家设定快捷键：1是10秒，2是20秒，以此类推。",
            "需加热物品放入后关门，按对应数字即可快捷开始加热，无需再按开始键。",
        ]),
        ("保温柜", "用于保温成品餐品和配好的饭底", "image6.png", [
            "保温柜需要购买最高温度在50度以上的。",
        ]),
    ]
    for name, use, image, notes in equipment:
        add_equipment_card(
            doc,
            name,
            use,
            image,
            notes,
            image_width=3.2 if image == "image6.png" else 6.8,
            image_height=5.9 if image == "image6.png" else 4.5,
        )

    doc.add_heading("低温机与冰箱", level=3)
    add_key_table(doc, [
        ("低温机", "用于煮鸡胸和温泉蛋。温度设置为65度，加热时间从到达65度时开始算；加水必须超过最低水位线。"),
        ("冷藏冰箱", "设定温度为0-2度之间。先设1度，再根据蔬菜结冰或食材变质情况微调。"),
        ("冷冻冰箱", "设定温度-18度，并确保实际温度达到-15度以下。前期要检查货物是否完全冷冻。"),
    ])

    doc.add_heading("设备日常保养", level=2)
    add_bullets(doc, [
        "烤炉：每日闭店清洗外观、底部水盘和烤架，加热管不可以直接洗刷。",
        "烤箱：每周进行内部清洁，风扇处避免长时间积灰；背面不要完全挡住排气孔。",
        "电饭锅：每周把锅盖背面的背板取下进行清洁，避免米饭残渣和蒸汽长期积累。",
        "低温机：每天用完拧开外罩冲洗干净，每周刷洗一次；有水垢时用醋煮泡并刷洗。",
        "冰箱：冷冻冰箱每周除霜，密封条擦洗防霉；冷藏冰箱如有积水及时清理。",
    ])
    doc.add_heading("常见简易故障", level=2)
    add_bullets(doc, [
        "烤炉：长时间未使用的加热管开启后冒白烟，多为残留油脂产生，烧热一会儿会消失。",
        "烤箱：食物上出现黑点时，通常需要彻底清理内部黑色油渍污渍和风扇。",
        "低温机：建议每家店准备两个，避免高峰期设备问题导致出餐断档。",
        "低温机：水位低于最低水位线无法启动；水垢积满会影响加热。",
        "低温机：如有异响或水流异常，应立刻停止使用并排查或更换机器。",
    ])


def chapter_workflow(doc):
    new_page(doc)
    doc.add_heading("二、全天工作流程", level=1)
    doc.add_heading("开早", level=2)
    doc.add_heading("早班岗位 8-7（根据单量分配1-2人）", level=3)
    add_numbers(doc, [
        "到店后第一件事：低温机加水煮温泉蛋，打开烤箱预热，打开炉灶烧两锅水。",
        "准备烤箱食材并摆入烤盘；需要炒和煮的蔬菜放台面备用。",
        "把所有米饭打出，做上第二轮米饭。",
        "四个炉头全部开启：炒洋葱、炒蘑菇、煮玉米粒、煮西兰花；随后继续炒洋葱和蘑菇，同时开始煎樟树椒、煮土豆、虾仁、全熟鸡蛋等。",
        "烤箱预热完成后，把食材放入并设定好时间。",
    ])
    doc.add_heading("晚班岗位 9-8（比早班晚1小时）", level=3)
    add_numbers(doc, [
        "到店后开始打饭底，放保温柜备用。",
        "打出第二轮米饭，视情况确定是否继续煮米饭。",
        "烤制牛五花，制备预估到中午12点的牛五花饭，放保温柜备用。",
        "制备预估到中午12点的三文鱼饭、鸡肉饭、鸡腿饭，放保温柜备用。",
        "制备少量牛排饭和波奇饭饭底，放保温柜备用。",
        "做土豆泥，盛装一部分备用，其余放凉后冷藏。",
        "整理打包包材，准备打包小菜。",
    ])
    add_para(doc, "预制作原则：数量最多预估到12点30分的使用量，不可多备。12点后尽量保证菜品新鲜。加热柜中未用完的餐品，需要确认西兰花和菠菜是否变黄，变黄不可使用。", bold_lead=True)

    doc.add_heading("午餐高峰", level=2)
    add_key_table(doc, [
        ("烤炉岗位", "负责现制牛排类餐品烤制，嫩肩和藤椒牛肉制作，以及配菜补充。"),
        ("配餐岗位", "接单后将已有成品交给打包岗；需要现制的单子告知烤炉岗，并准备对应饭底。同步掌握成品、半成品饭底和配菜数量。"),
        ("打包岗位", "现成餐品直接打包，穿插打包现做好的餐品。"),
    ])

    doc.add_heading("下午备货", level=2)
    add_bullets(doc, [
        "蔬菜类：切洋葱、切白玉菇、切西兰花、圣女果、切胡萝卜条、土豆削皮。",
        "半成品制作：煮裙带菜冷藏备用，煮菠菜冷藏/冷冻备用，煮鸡胸放凉后冷藏备用。",
        "小料类：蒜香黄油制备，酱汁、小菜提前装盛，冷藏备用。",
        "下午4点后开始准备晚餐时段原料和配菜。晚餐来单较分散，不需要制备大量饭底和成品。",
        "土豆泥、手枪腿等制作时间长的产品，6点后卖完可以沽清，不再制作。",
    ])

    doc.add_heading("打烊", level=2)
    add_bullets(doc, [
        "食材冷藏：蔬菜报废处理；鸡蛋、肉类未超保质期的，用保鲜盒收纳后放入冰箱。",
        "增添鸡蛋：放入第二天早上需要煮的鸡蛋。",
        "肉类解冻：清点剩余肉类数量，把第二天需要用的量拿出来解冻。",
        "米饭预约：确保全部预约成功，不要错误设置成立即煮饭。",
        "卫生清扫：地面和台面清洁干净。",
        "设备清洁：烤炉外观和水盘、锅具厨具、分数盒清洗干净。",
        "确认切断电源、燃气，并拍照留存。",
    ])


def chapter_production(doc):
    new_page(doc)
    doc.add_heading("三、餐品制作和打包", level=1)
    doc.add_paragraph("本章为打印清晰版：配方、制备方式和出餐配比优先保证可读；长图按菜品完整边界切开，不切断文字。")

    add_packed_item_images(doc, "image7.png", "蔬菜原料制备", boundaries=[
        0, 550, 958, 1450, 1874, 2298, 3530, 3938, 4466,
        4974, 5330, 5706, 6146, 6418, 6774, 6910, 7198, 7350, 7444,
    ], width_cm=18.0, max_display_height_cm=20.0)
    new_page(doc)

    doc.add_heading("部分原料制作配比", level=2)
    add_figure(doc, "image8.png", "部分原料制作中需要的配比", width_cm=15.8)
    add_figure(doc, "image9.png", "烤箱预热到200度后的加工时间", width_cm=6.8)
    new_page(doc)

    doc.add_heading("饭底制备", level=2)
    add_figure(doc, "image10.png", "米饭打底，蔬菜铺上层，按图示配比制备", width_cm=9.3)
    new_page(doc)

    add_packed_item_images(doc, "image11.png", "肉类制备", boundaries=[
        0, 654, 1442, 1902, 2534, 3354, 3934, 4608,
    ], width_cm=18.0, max_display_height_cm=20.0)
    new_page(doc)

    doc.add_heading("小菜、酱汁与调味品", level=2)
    add_figure(doc, "image12.png", "小菜与酱汁制备表", width_cm=15.8)
    add_figure(doc, "image13.png", "调味品保存方式", width_cm=9.8)
    new_page(doc)

    doc.add_heading("出餐成品配比", level=2)
    add_figure(doc, "image14.png", "牛排饭出餐成品配比", width_cm=15.8)
    add_figure(doc, "image15.png", "波奇饭出餐成品配比", width_cm=15.8)
    add_figure(doc, "image16.png", "土豆泥牛排出餐成品配比", width_cm=15.8)


def chapter_food_safety(doc):
    new_page(doc)
    doc.add_heading("四、食品安全", level=1)
    doc.add_heading("夏季食材存储温度管理", level=2)
    add_numbers(doc, [
        "危险温度带：10度-60度。热制好的配菜常温存放不得超过4小时。",
        "蔬菜类（洋葱、白玉菇、胡萝卜、西兰花、拌好的菠菜）不隔夜，当日报废。",
        "已经做熟未使用的原料（如煮好的鸡胸、裙带菜），持续冷藏状态下不超过36小时。",
        "已经做熟并且已使用的原料（如开袋鸡胸、烤好的鸡腿和三文鱼），冷藏状态下不超过24小时；如果常温存放较长，当天内用完，不可隔夜。",
        "打好的饭底在保温柜热藏不超过3小时，超时需进冷藏存放，总时长不超过5小时。",
    ])
    add_para(doc, "例：10点打好的饭底，13点仍未销售完，需从柜中取出放入冷藏，15点前报废，不得用于晚餐时段。")

    doc.add_heading("人员卫生管理", level=2)
    add_bullets(doc, [
        "进厨房立刻戴网帽，不得裸露头发。",
        "禁止厨房内吸烟、饮酒、躺卧。",
        "不得用手直接接触食材，必须戴手套；一次性手套及时更换。",
        "保持个人卫生，围裙定期清洗，不得佩戴首饰，指甲修剪干净。",
        "使用卫生间、咳嗽打喷嚏擦鼻涕、收垃圾废弃物、清扫垃圾后需要洗手。",
        "个人物品集中有序存放，不得和食材混放。",
    ])
    doc.add_heading("门店卫生管理", level=2)
    add_bullets(doc, [
        "食材和包材不可以直接落地。",
        "食材不和垃圾桶混放。",
        "高峰期后及时清理地面和台面，保持洁净。",
        "清洗拖把不得和食材混在一起。",
        "隔油池每周清理一次。",
        "一次性手套及时更换，不要在不同食材之中搅动。",
        "掉在地上的成品食材不得二次使用。",
        "拆装订书针时，必须距离餐品50cm以上。",
        "不要在锅边拆袋，避免塑料袋掉进锅内。",
        "切肉时注意手套是否完整，避免把手套切碎带进餐品。",
        "出餐时保证餐盒四周干净，切肉手套不要直接接触餐盒；打包岗位检查餐盒清洁度。",
    ])
    doc.add_heading("冰箱管理", level=2)
    add_bullets(doc, [
        "冰箱内食材不得带外包装箱，需用保鲜盒或收纳筐。",
        "冰箱内不得存放除店用食材外的其他食品。",
        "门店和冰箱内不得出现无标签、标签过保质期的产品。",
        "未封膜加盖的食品不得直接放进冰箱。",
        "冰箱内生熟不得混放。",
        "不可以使用出餐包材存储半成品后放入冰箱。",
        "冰箱定期清霜清洁，随手关门，晚上拍照确认关好。",
        "冰箱货物执行先进先出，尤其是冷藏食材。",
    ])
    doc.add_heading("虫害管理", level=2)
    add_bullets(doc, [
        "日常保证卫生死角无原料或大量食物残渣残留。",
        "门店与正规消杀公司签订年度消杀合同，确保每月按时处理。",
        "独立门店需有吊顶，避免鼠患发生。",
        "出餐过程中，不使用的食材尽快封膜加盖，防止虫害进入。",
    ])


def chapter_close_and_exception(doc):
    new_page(doc)
    doc.add_heading("五、每日闭店拍照", level=1)
    add_numbers(doc, [
        "需要报废的食材（称重拍照）。",
        "留存食材已放入冰箱，防止有冷藏保存食材遗漏在常温下。",
        "微波炉中无遗漏餐食或原料。",
        "烤炉、电磁炉、加热柜关闭；电的设备拍空开关闭，燃气拍燃气阀门。",
        "第二天解冻的肉。",
        "冰箱关紧，无漏缝。",
        "整店清洁后的照片。",
        "检查平台沽清商品是否上架。",
    ])
    doc.add_heading("六、异常订单处理方式", level=1)
    add_figure(doc, "image17.jpeg", "异常订单处理逻辑图", width_cm=15.8)


def get_complaint_table_rows():
    from docx import Document as SourceDocument
    source = SourceDocument(str(SRC))
    if not source.tables:
        return []
    rows = []
    for row in source.tables[0].rows[1:]:
        cells = []
        for cell in row.cells:
            text = "\n".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            cells.append(text)
        if any(cells):
            rows.append(cells)
    return rows


def chapter_complaints(doc):
    new_page(doc)
    doc.add_heading("七、客诉参考处理方式", level=1)
    doc.add_heading("客诉处理黄金公式", level=2)
    add_para(doc, "向客户道歉 + 向客户解释问题产生的原因 + 解决方案")
    add_key_table(doc, [
        ("第一步", "不管问题出在哪里，第一时间先回复顾客，先道歉并表达态度。"),
        ("第二步", "解释问题产生原因。客观事实马上认错；主观问题不做绝对定责，但要承诺追责和改进。"),
        ("第三步", "及时给出解决方案，包括部分退款、补送、下次赠送小菜、整单退款等。"),
    ])
    doc.add_heading("在线沟通和电话原则", level=2)
    add_bullets(doc, [
        "核心目标是不要差评，在此基础上再考虑其他。",
        "无论是我们错还是用户找茬，第一件事是安抚情绪：先承认、道歉、表态，再查证和解决。",
        "对用户做判断：有价值的老用户真诚解决；找茬用户重点控制差评，不要激怒对方。",
        "问题先定级，再匹配补救措施。",
    ])

    rows = get_complaint_table_rows()
    by_type = {}
    for cells in rows:
        by_type.setdefault(cells[0], []).append(cells)

    for problem_type, items in by_type.items():
        doc.add_heading(problem_type, level=2)
        for cells in items:
            _, common, strategy, remedy, script = (cells + [""] * 5)[:5]
            doc.add_heading(common, level=3)
            add_key_table(doc, [
                ("回答策略", strategy.replace("\n", "；")),
                ("补救措施", remedy.replace("\n", "；")),
                ("话术参考", script.replace("\n", "；")),
            ], widths=(2.7, 12.5))


def patch_docx_fonts(path):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        with zipfile.ZipFile(path, "r") as zin:
            zin.extractall(tmp_path)

        for rel in ["word/document.xml", "word/styles.xml"]:
            xml_path = tmp_path / rel
            if not xml_path.exists():
                continue
            tree = etree.parse(str(xml_path))
            root = tree.getroot()
            for rfonts in root.xpath(".//w:rFonts", namespaces=ns):
                for attr in ["ascii", "hAnsi", "eastAsia", "cs"]:
                    rfonts.set(qn(f"w:{attr}"), FONT)
            # Make fields update if the reader supports it.
            settings = tmp_path / "word/settings.xml"
            if settings.exists():
                st = etree.parse(str(settings))
                sr = st.getroot()
                if not sr.xpath(".//w:updateFields", namespaces=ns):
                    upd = OxmlElement("w:updateFields")
                    upd.set(qn("w:val"), "true")
                    sr.append(upd)
                    st.write(str(settings), encoding="UTF-8", xml_declaration=True, standalone=True)
            tree.write(str(xml_path), encoding="UTF-8", xml_declaration=True, standalone=True)

        tmp_out = path.with_suffix(".tmp.docx")
        with zipfile.ZipFile(tmp_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for file in tmp_path.rglob("*"):
                if file.is_file():
                    zout.write(file, file.relative_to(tmp_path))
        tmp_out.replace(path)


def main():
    extract_media()
    doc = Document()
    setup_document(doc)
    add_cover(doc)
    add_contents(doc)
    chapter_equipment(doc)
    chapter_workflow(doc)
    chapter_production(doc)
    chapter_food_safety(doc)
    chapter_close_and_exception(doc)
    chapter_complaints(doc)
    doc.save(OUT)
    patch_docx_fonts(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
